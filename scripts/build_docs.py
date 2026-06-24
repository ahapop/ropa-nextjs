# -*- coding: utf-8 -*-
"""แปลง Markdown (subset) -> .docx ฟอนต์ไทย  ·  python build_docs.py <in.md> <out.docx> "<title>" "<subtitle>" """
import sys, re, datetime, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "TH Sarabun New"
BODY = 15
H = {1: 22, 2: 18, 3: 16, 4: 15}
NAVY = RGBColor(0x16, 0x44, 0x7F)
MUTED = RGBColor(0x6b, 0x77, 0x85)
LOGO = "public/btsc-logo-transparent.png"

def set_cs(rpr, name):
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), name)

def style_font(style, name, size, bold=False, color=None):
    style.font.name = name; style.font.size = Pt(size); style.font.bold = bold
    if color is not None: style.font.color.rgb = color
    set_cs(style.element.get_or_add_rPr(), name)

def runset(run, name=FONT):
    run.font.name = name
    set_cs(run._element.get_or_add_rPr(), name)

def _field(p, instr, placeholder="", size=None, color=None):
    """แทรก Word field (PAGE / NUMPAGES / TOC) ลงใน paragraph"""
    r = p.add_run(); rel = r._r
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); rel.append(b)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr; rel.append(it)
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate'); rel.append(sep)
    if placeholder:
        t = OxmlElement('w:t'); t.text = placeholder; rel.append(t)
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); rel.append(e)
    runset(r)
    if size: r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    return r

def add_toc(doc):
    p = doc.add_paragraph()
    _field(p, 'TOC \\o "1-3" \\h \\z \\u',
           placeholder="(สารบัญจะสร้างอัตโนมัติเมื่อเปิดไฟล์ · หากไม่ขึ้น คลิกขวาที่นี่ > Update Field)")

def add_footer(doc):
    fp = doc.sections[0].footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = fp.add_run("หน้า "); runset(r0); r0.font.size = Pt(12); r0.font.color.rgb = MUTED
    _field(fp, "PAGE", "1", size=12, color=MUTED)
    r1 = fp.add_run(" / "); runset(r1); r1.font.size = Pt(12); r1.font.color.rgb = MUTED
    _field(fp, "NUMPAGES", "1", size=12, color=MUTED)

def set_update_fields(doc):
    el = doc.settings.element
    uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); el.append(uf)

def emphasize(p, text):
    """แตก **bold** และ `code` เป็น runs"""
    for part in re.split(r'(\*\*.+?\*\*|`.+?`)', text):
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True; runset(r)
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); runset(r, 'Consolas')
        else:
            r = p.add_run(part); runset(r)

def is_table_sep(s):
    return bool(re.fullmatch(r'\s*\|?[\s:\-|]+\|?\s*', s)) and '-' in s

def cells(row):
    row = row.strip()
    if row.startswith('|'): row = row[1:]
    if row.endswith('|'): row = row[:-1]
    return [c.strip() for c in row.split('|')]

def add_table(doc, rows):
    hdr = cells(rows[0]); body = [cells(r) for r in rows[2:]]
    ncol = max([len(hdr)] + [len(r) for r in body])
    t = doc.add_table(rows=1, cols=ncol); t.style = 'Light Grid Accent 1'
    for j in range(ncol):
        c = t.rows[0].cells[j]; c.paragraphs[0].text = ''
        emphasize(c.paragraphs[0], hdr[j] if j < len(hdr) else '')
        for run in c.paragraphs[0].runs: run.bold = True
    for r in body:
        cs = t.add_row().cells
        for j in range(ncol):
            cs[j].paragraphs[0].text = ''
            emphasize(cs[j].paragraphs[0], r[j] if j < len(r) else '')
    doc.add_paragraph()

def build(md, out, title, subtitle):
    doc = Document()
    style_font(doc.styles['Normal'], FONT, BODY)
    for i in (1, 2, 3, 4):
        try: style_font(doc.styles[f'Heading {i}'], FONT, H[i], bold=True, color=NAVY)
        except KeyError: pass
    for s in ('List Bullet', 'List Number'):
        try: style_font(doc.styles[s], FONT, BODY)
        except KeyError: pass

    set_update_fields(doc)   # ให้ Word เสนออัปเดต field (สารบัญ/เลขหน้า) ตอนเปิด
    add_footer(doc)          # เลขหน้าใน footer

    # ปก — โลโก้
    if os.path.exists(LOGO):
        lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(LOGO, width=Inches(1.7))
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(title); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY; runset(r)
    if subtitle:
        sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sp.add_run(subtitle); r.font.size = Pt(15); r.font.color.rgb = RGBColor(0x6b,0x77,0x85); runset(r)
    dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp.add_run("จัดทำเมื่อ " + datetime.date.today().strftime("%d/%m/%Y")); r.font.size = Pt(13); runset(r)
    doc.add_page_break()

    lines = md.split('\n'); i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        s = line.strip()
        # table
        if s.startswith('|') and i + 1 < len(lines) and is_table_sep(lines[i+1]):
            blk = [line, lines[i+1]]; i += 2
            while i < len(lines) and lines[i].strip().startswith('|'):
                blk.append(lines[i]); i += 1
            add_table(doc, blk); continue
        # code fence
        if s.startswith('```'):
            i += 1; code = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run('\n'.join(code)); r.font.size = Pt(11); runset(r, 'Consolas')
            continue
        if s == '[[PAGEBREAK]]':
            doc.add_page_break(); i += 1; continue
        if s == '[[TOC]]':
            add_toc(doc); i += 1; continue
        m = re.match(r'!\[(.*?)\]\((.*?)\)', s)
        if m:
            cap, path = m.group(1), m.group(2)
            try:
                doc.add_picture(path, width=Inches(6.4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if cap:
                    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rr = cp.add_run("▲ " + cap); rr.italic = True; rr.font.size = Pt(12)
                    rr.font.color.rgb = RGBColor(0x6b,0x77,0x85); runset(rr)
            except Exception as ex:
                doc.add_paragraph("[ไม่พบรูป: " + path + "]")
            i += 1; continue
        m = re.match(r'(#{1,4})\s+(.*)', s)
        if m:
            lvl = len(m.group(1)); h = doc.add_heading(level=lvl)
            emphasize(h, m.group(2)); i += 1; continue
        if s in ('---', '***', '___'):
            doc.add_paragraph('—' * 30); i += 1; continue
        m = re.match(r'[-*]\s+(.*)', s)
        if m:
            p = doc.add_paragraph(style='List Bullet'); emphasize(p, m.group(1)); i += 1; continue
        m = re.match(r'\d+\.\s+(.*)', s)
        if m:
            p = doc.add_paragraph(style='List Number'); emphasize(p, m.group(1)); i += 1; continue
        if s.startswith('> '):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18)
            emphasize(p, s[2:]); i += 1; continue
        if s == '':
            i += 1; continue
        p = doc.add_paragraph(); emphasize(p, s); i += 1

    doc.save(out)
    print("saved", out)

if __name__ == '__main__':
    in_md, out_docx, title = sys.argv[1], sys.argv[2], sys.argv[3]
    subtitle = sys.argv[4] if len(sys.argv) > 4 else ""
    with open(in_md, encoding='utf-8') as f:
        build(f.read(), out_docx, title, subtitle)
